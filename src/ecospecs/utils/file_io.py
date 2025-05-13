from pathlib import Path
from docx import Document
import pdfplumber
from typing import List

def save_tables(tables: List[List[List[str]]], output_path: Path, file_format: str = "docx"):
    """Save tables to DOCX/PDF file in output directory"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    if file_format == "docx":
        _save_as_docx(tables, output_path)
    elif file_format == "pdf":
        _save_as_pdf(tables, output_path)
    else:
        raise ValueError(f"Unsupported format: {file_format}")

def _save_as_docx(tables: List[List[List[str]]], output_path: Path):
    doc = Document()
    for table_data in tables:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
        doc.add_paragraph()  # Add spacing between tables
    doc.save(output_path)

# def _save_as_pdf(tables: List[List[List[str]]], output_path: Path):
#     with pdfplumber.open() as pdf:
#         # PDF creation is more complex - this is a simplified version
#         # Consider using reportlab for better PDF generation
#         pdf.save(output_path)